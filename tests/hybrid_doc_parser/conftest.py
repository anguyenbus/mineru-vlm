# conftest.py — shared pytest fixtures for hybrid_doc_parser test suite.

from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture(scope="session")
def docx_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Programmatically generated .docx fixture — not committed as binary.

    Creates a minimal DOCX file using python-docx at test-collection time and
    writes it to a session-scoped temporary directory. The file contains a
    heading and a paragraph, which is sufficient to exercise the Docling
    parsing pipeline in unit tests.

    Args:
        tmp_path_factory: pytest's session-scoped temporary directory factory.

    Returns:
        Path to the generated ``test.docx`` file.

    Raises:
        ImportError: When ``python-docx`` is not installed in the environment.
            Install with ``uv add --dev python-docx`` or
            ``pip install python-docx``.
    """
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for the docx_fixture. "
            "Install with: uv add --dev 'python-docx>=1.0'"
        ) from exc

    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("Hello from python-docx fixture.")
    path = tmp_path_factory.mktemp("docx") / "test.docx"
    doc.save(str(path))
    return path
